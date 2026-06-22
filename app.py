import streamlit as st
import whisper
from pydub import AudioSegment
from docx import Document
from datetime import datetime
import torch
import tempfile
import os

st.set_page_config(page_title="會議紀要轉化器", layout="centered")
st.title("🎙️ 自動會議紀要與待辦轉化器")
st.markdown("**支援音檔 / 影片 • 粵語 / 普通話 / 英語**")

with st.sidebar:
    st.header("設定")
    model_name = st.selectbox(
        "選擇 Whisper 模型",
        options=["medium", "large-v3", "small"],
        index=0,
        help="large-v3 最準確，但較慢且耗記憶體"
    )

uploaded_file = st.file_uploader(
    "上傳會議音檔或影片",
    type=["mp3", "wav", "m4a", "mp4", "mov", "avi", "mkv"],
    help="支援常見格式"
)

if uploaded_file:
    st.success(f"✅ 已上傳：{uploaded_file.name}")
    
    if st.button("🚀 開始轉錄", type="primary"):
        with st.spinner("正在轉錄中... 請耐心等待"):
            try:
                # 儲存上傳檔案
                with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name) as tmp:
                    tmp.write(uploaded_file.getbuffer())
                    temp_path = tmp.name

                # 轉成 wav
                audio = AudioSegment.from_file(temp_path)
                wav_path = temp_path + ".wav"
                audio.export(wav_path, format="wav")

                # 載入模型（快取加速）
                @st.cache_resource
                def load_model(name):
                    return whisper.load_model(name)
                
                model = load_model(model_name)

                result = model.transcribe(wav_path, language=None)
                transcript = result["text"]

                st.subheader("📝 轉錄結果")
                st.text_area("完整逐字稿", transcript, height=400)

                # 生成 Word
                doc = Document()
                doc.add_heading('會議紀要與待辦事項', 0)
                doc.add_paragraph(f'日期：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
                doc.add_paragraph(f'檔案：{uploaded_file.name}')
                doc.add_paragraph('='*50)
                doc.add_heading('📋 完整逐字稿', level=1)
                doc.add_paragraph(transcript)
                
                doc.add_page_break()
                doc.add_heading('✅ 待辦事項', level=1)
                doc.add_paragraph("請手動整理或使用 LLM 提取")

                docx_name = f"會議紀要_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                doc.save(docx_name)

                # 下載按鈕
                col1, col2 = st.columns(2)
                with col1:
                    with open(docx_name, "rb") as f:
                        st.download_button("📥 下載 Word 會議紀要", f, file_name=docx_name)
                with col2:
                    st.download_button("📥 下載 TXT 逐字稿", transcript, 
                                     file_name=f"逐字稿_{datetime.now().strftime('%Y%m%d_%H%M')}.txt")

            except Exception as e:
                st.error(f"錯誤：{str(e)}")
            finally:
                # 清理暫存檔
                for p in [temp_path, wav_path, docx_name]:
                    if os.path.exists(p):
                        os.remove(p)
